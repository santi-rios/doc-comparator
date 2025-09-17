from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def create_dummy_files():
    with open('/home/santi/doc-comparator/original.txt', 'r') as f:
        original_text = f.read()

    # Create DOCX
    doc = Document()
    doc.add_paragraph("This is the modified document. It contains some text that will be compared, but with some changes.")
    doc.add_paragraph(original_text)
    doc.save('/home/santi/doc-comparator/document.docx')

    # Create PDF
    c = canvas.Canvas('/home/santi/doc-comparator/document.pdf', pagesize=letter)
    width, height = letter
    text = c.beginText(50, height - 50)
    text.setFont("Helvetica", 12)
    text.textLine("This is the original document.")
    text.textLine(original_text)
    c.drawText(text)
    c.save()

if __name__ == '__main__':
    create_dummy_files()
